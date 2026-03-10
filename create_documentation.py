"""
Script to generate Project Documentation in Word format
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_with_number(doc, text, level=1):
    """Add a numbered heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def create_table_with_header(doc, headers, data, col_widths=None):
    """Create a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '4472C4')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for row_data in data:
        row = table.add_row()
        for i, cell_data in enumerate(row_data):
            row.cells[i].text = str(cell_data)
    
    return table

def create_documentation():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # ==================== TITLE PAGE ====================
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("HEALTHCARE DATA ANALYSIS")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Using Visualization and Predictive Models")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(18)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Project Type
    project_type = doc.add_paragraph()
    project_type.add_run("A Project Report").bold = True
    project_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    submitted = doc.add_paragraph()
    submitted.add_run("Submitted in partial fulfillment of the requirements")
    submitted.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    degree = doc.add_paragraph()
    degree.add_run("for the award of the degree of")
    degree.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    degree_name = doc.add_paragraph()
    degree_name.add_run("BACHELOR OF TECHNOLOGY").bold = True
    degree_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    branch = doc.add_paragraph()
    branch.add_run("in")
    branch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    branch_name = doc.add_paragraph()
    branch_name.add_run("COMPUTER SCIENCE AND ENGINEERING").bold = True
    branch_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Submitted By
    by = doc.add_paragraph()
    by.add_run("Submitted By:").bold = True
    by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    name = doc.add_paragraph()
    name.add_run("S V L JYOTHIKA NOOKALA")
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Year
    year = doc.add_paragraph()
    year.add_run("March 2026").bold = True
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ==================== CERTIFICATE ====================
    cert_title = doc.add_paragraph()
    cert_title.add_run("CERTIFICATE").bold = True
    cert_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cert_title.runs[0].font.size = Pt(16)
    
    doc.add_paragraph()
    
    cert_text = doc.add_paragraph()
    cert_text.add_run(
        "This is to certify that the project entitled "
    )
    cert_text.add_run('"Healthcare Data Analysis Using Visualization and Predictive Models"').bold = True
    cert_text.add_run(
        " is a bonafide work carried out by "
    )
    cert_text.add_run("S V L JYOTHIKA NOOKALA").bold = True
    cert_text.add_run(
        " in partial fulfillment of the requirements for the award of Bachelor of Technology "
        "in Computer Science and Engineering."
    )
    cert_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signatures
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.cell(0, 0).text = "Project Guide"
    sig_table.cell(0, 1).text = "Head of Department"
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_para.add_run("Date: _______________")
    
    place_para = doc.add_paragraph()
    place_para.add_run("Place: _______________")
    
    doc.add_page_break()
    
    # ==================== DECLARATION ====================
    decl_title = doc.add_paragraph()
    decl_title.add_run("DECLARATION").bold = True
    decl_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    decl_title.runs[0].font.size = Pt(16)
    
    doc.add_paragraph()
    
    decl_text = doc.add_paragraph()
    decl_text.add_run(
        "I hereby declare that the project entitled "
    )
    decl_text.add_run('"Healthcare Data Analysis Using Visualization and Predictive Models"').bold = True
    decl_text.add_run(
        " submitted for the Bachelor of Technology degree in Computer Science and Engineering "
        "is my original work and the project has not formed the basis for the award of any other "
        "degree, diploma, fellowship or any other similar title."
    )
    decl_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    sig = doc.add_paragraph()
    sig.add_run("S V L JYOTHIKA NOOKALA")
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()
    
    place_date = doc.add_paragraph()
    place_date.add_run("Place: _______________")
    
    date_decl = doc.add_paragraph()
    date_decl.add_run("Date: _______________")
    
    doc.add_page_break()
    
    # ==================== ACKNOWLEDGEMENT ====================
    ack_title = doc.add_paragraph()
    ack_title.add_run("ACKNOWLEDGEMENT").bold = True
    ack_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ack_title.runs[0].font.size = Pt(16)
    
    doc.add_paragraph()
    
    ack_text = doc.add_paragraph()
    ack_text.add_run(
        "I would like to express my sincere gratitude to all who helped me in completing "
        "this project successfully. I am thankful to my project guide for providing valuable "
        "guidance and constant encouragement throughout the project work.\n\n"
        "I am thankful to the Head of the Department and all faculty members for their support "
        "and cooperation. I also extend my thanks to my family and friends for their moral support "
        "and encouragement during this project.\n\n"
        "Finally, I thank all those who directly or indirectly helped me in the successful "
        "completion of this project."
    )
    ack_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    name_ack = doc.add_paragraph()
    name_ack.add_run("S V L JYOTHIKA NOOKALA")
    name_ack.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()
    
    # ==================== ABSTRACT ====================
    abs_title = doc.add_paragraph()
    abs_title.add_run("ABSTRACT").bold = True
    abs_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    abs_title.runs[0].font.size = Pt(16)
    
    doc.add_paragraph()
    
    abstract = doc.add_paragraph()
    abstract.add_run(
        "With the rapid growth of digital healthcare systems, large volumes of medical data are "
        "continuously generated from patient records, clinical reports, laboratory results, imaging "
        "studies, and health monitoring devices. Data visualization and predictive analytics play a "
        "vital role in transforming these datasets into meaningful insights that support better "
        "understanding, early risk detection, and proactive intervention strategies.\n\n"
        "This project focuses on systematically analyzing healthcare datasets and presenting the "
        "results through interactive visual dashboards integrated with predictive models. The system "
        "uses Machine Learning algorithms including Random Forest Classifier for Heart Disease prediction "
        "and Gradient Boosting Classifier for Diabetes Risk assessment.\n\n"
        "The application provides a comprehensive platform for healthcare data analysis featuring "
        "multi-disease prediction, health score calculation (0-100 scale), BMI calculator, analytics "
        "dashboard with rich visualizations, prediction history tracking, and CSV/PDF export capabilities. "
        "Security features include CSRF protection, rate-limited login, session timeout, strong password "
        "enforcement, and bcrypt encryption.\n\n"
        "The system achieves approximately 81% accuracy for heart disease prediction and 75% accuracy "
        "for diabetes risk assessment using 5-fold cross-validation on augmented datasets."
    )
    abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    keywords = doc.add_paragraph()
    keywords.add_run("Keywords: ").bold = True
    keywords.add_run(
        "Healthcare Analytics, Machine Learning, Predictive Models, Data Visualization, "
        "Heart Disease Prediction, Diabetes Risk Assessment, Flask, MongoDB, Random Forest, "
        "Gradient Boosting"
    )
    
    doc.add_page_break()
    
    # ==================== TABLE OF CONTENTS ====================
    toc_title = doc.add_paragraph()
    toc_title.add_run("TABLE OF CONTENTS").bold = True
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.runs[0].font.size = Pt(16)
    
    doc.add_paragraph()
    
    toc_items = [
        ("Certificate", "ii"),
        ("Declaration", "iii"),
        ("Acknowledgement", "iv"),
        ("Abstract", "v"),
        ("Table of Contents", "vi"),
        ("List of Tables", "viii"),
        ("List of Figures", "ix"),
        ("", ""),
        ("CHAPTER 1: INTRODUCTION", "1"),
        ("    1.1 Overview", "1"),
        ("    1.2 Problem Statement", "2"),
        ("    1.3 Objectives", "2"),
        ("    1.4 Scope of the Project", "3"),
        ("", ""),
        ("CHAPTER 2: LITERATURE REVIEW", "4"),
        ("    2.1 Existing Systems", "4"),
        ("    2.2 Proposed System", "5"),
        ("    2.3 Advantages of Proposed System", "6"),
        ("", ""),
        ("CHAPTER 3: SYSTEM ANALYSIS", "7"),
        ("    3.1 Feasibility Study", "7"),
        ("    3.2 System Requirements", "8"),
        ("    3.3 Software Requirements", "9"),
        ("    3.4 Hardware Requirements", "9"),
        ("", ""),
        ("CHAPTER 4: SYSTEM DESIGN", "10"),
        ("    4.1 System Architecture", "10"),
        ("    4.2 Data Flow Diagrams", "11"),
        ("    4.3 ER Diagram", "13"),
        ("    4.4 Database Design", "14"),
        ("", ""),
        ("CHAPTER 5: IMPLEMENTATION", "16"),
        ("    5.1 Technology Stack", "16"),
        ("    5.2 Machine Learning Models", "17"),
        ("    5.3 Module Description", "19"),
        ("    5.4 Code Implementation", "21"),
        ("", ""),
        ("CHAPTER 6: TESTING", "25"),
        ("    6.1 Testing Methodology", "25"),
        ("    6.2 Test Cases", "26"),
        ("    6.3 Results", "27"),
        ("", ""),
        ("CHAPTER 7: RESULTS AND DISCUSSION", "28"),
        ("    7.1 Screenshots", "28"),
        ("    7.2 Model Performance", "32"),
        ("", ""),
        ("CHAPTER 8: CONCLUSION AND FUTURE SCOPE", "34"),
        ("    8.1 Conclusion", "34"),
        ("    8.2 Future Enhancements", "35"),
        ("", ""),
        ("REFERENCES", "36"),
        ("APPENDIX", "37"),
    ]
    
    for item, page in toc_items:
        if item:
            toc_para = doc.add_paragraph()
            toc_para.add_run(item)
            if page:
                toc_para.add_run("\t" * 5 + page)
    
    doc.add_page_break()
    
    # ==================== LIST OF TABLES ====================
    lot_title = doc.add_paragraph()
    lot_title.add_run("LIST OF TABLES").bold = True
    lot_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lot_title.runs[0].font.size = Pt(16)
    
    doc.add_paragraph()
    
    tables_list = [
        ("Table 3.1", "Software Requirements", "9"),
        ("Table 3.2", "Hardware Requirements", "9"),
        ("Table 4.1", "Users Collection Schema", "14"),
        ("Table 4.2", "Predictions Collection Schema", "15"),
        ("Table 5.1", "Technology Stack", "16"),
        ("Table 5.2", "Heart Disease Model Specifications", "17"),
        ("Table 5.3", "Diabetes Model Specifications", "18"),
        ("Table 5.4", "Input Features Description", "18"),
        ("Table 5.5", "API Endpoints", "20"),
        ("Table 6.1", "Test Cases", "26"),
        ("Table 7.1", "Model Performance Metrics", "32"),
    ]
    
    for num, title, page in tables_list:
        table_para = doc.add_paragraph()
        table_para.add_run(f"{num}: {title}")
        table_para.add_run("\t" * 4 + page)
    
    doc.add_page_break()
    
    # ==================== LIST OF FIGURES ====================
    lof_title = doc.add_paragraph()
    lof_title.add_run("LIST OF FIGURES").bold = True
    lof_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lof_title.runs[0].font.size = Pt(16)
    
    doc.add_paragraph()
    
    figures_list = [
        ("Figure 4.1", "System Architecture Diagram", "10"),
        ("Figure 4.2", "Context Level DFD (Level 0)", "11"),
        ("Figure 4.3", "Level 1 DFD", "12"),
        ("Figure 4.4", "ER Diagram", "13"),
        ("Figure 7.1", "Home Page", "28"),
        ("Figure 7.2", "Registration Page", "28"),
        ("Figure 7.3", "Login Page", "29"),
        ("Figure 7.4", "Dashboard", "29"),
        ("Figure 7.5", "Health Assessment Form", "30"),
        ("Figure 7.6", "Prediction Result Page", "30"),
        ("Figure 7.7", "Analytics Dashboard", "31"),
        ("Figure 7.8", "Prediction History", "31"),
        ("Figure 7.9", "Model Accuracy Comparison", "32"),
    ]
    
    for num, title, page in figures_list:
        fig_para = doc.add_paragraph()
        fig_para.add_run(f"{num}: {title}")
        fig_para.add_run("\t" * 4 + page)
    
    doc.add_page_break()
    
    # ==================== CHAPTER 1: INTRODUCTION ====================
    ch1_title = doc.add_heading("CHAPTER 1: INTRODUCTION", level=1)
    
    doc.add_heading("1.1 Overview", level=2)
    
    overview = doc.add_paragraph()
    overview.add_run(
        "Healthcare data analysis has emerged as one of the most critical applications of data science "
        "and machine learning in recent years. With the exponential growth of electronic health records (EHRs), "
        "medical imaging data, and patient monitoring systems, healthcare organizations are sitting on vast "
        "amounts of data that can be leveraged for improved patient outcomes.\n\n"
        "This project presents a comprehensive web-based healthcare data analysis platform that combines "
        "data visualization techniques with predictive machine learning models. The system is designed to "
        "analyze patient health metrics and provide risk assessments for two major health conditions:\n\n"
        "• Heart Disease - A leading cause of mortality worldwide\n"
        "• Diabetes - A chronic metabolic disorder affecting millions\n\n"
        "The application uses Random Forest and Gradient Boosting algorithms, which are well-suited for "
        "medical diagnosis tasks due to their ability to handle complex, non-linear relationships in data "
        "while maintaining interpretability."
    )
    overview.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading("1.2 Problem Statement", level=2)
    
    problem = doc.add_paragraph()
    problem.add_run(
        "Early detection of diseases like heart disease and diabetes is crucial for effective treatment "
        "and management. However, manual analysis of multiple health parameters is time-consuming and "
        "prone to human error. Traditional diagnostic methods often miss subtle patterns that could "
        "indicate early-stage disease risk.\n\n"
        "Key challenges addressed by this project:\n\n"
        "• Lack of accessible, user-friendly health risk assessment tools\n"
        "• Difficulty in tracking health trends over time\n"
        "• Limited integration of multiple health parameters for comprehensive analysis\n"
        "• Need for data visualization to understand health patterns\n"
        "• Absence of personalized health recommendations based on individual data"
    )
    problem.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading("1.3 Objectives", level=2)
    
    obj = doc.add_paragraph()
    obj.add_run("The primary objectives of this project are:\n\n")
    obj.add_run(
        "1. To develop a machine learning-based system for predicting heart disease and diabetes risk\n\n"
        "2. To create an interactive web application with user-friendly interfaces for health data input\n\n"
        "3. To implement data visualization dashboards for analyzing health trends\n\n"
        "4. To provide personalized health tips based on prediction results\n\n"
        "5. To maintain a secure system with proper authentication and data protection\n\n"
        "6. To enable users to track their health assessments over time\n\n"
        "7. To achieve high accuracy in disease risk prediction using ensemble learning methods"
    )
    obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading("1.4 Scope of the Project", level=2)
    
    scope = doc.add_paragraph()
    scope.add_run(
        "The scope of this project encompasses:\n\n"
        "• User registration and authentication system with security features\n"
        "• Multi-disease risk prediction (Heart Disease and Diabetes)\n"
        "• Health score calculation on a 0-100 scale\n"
        "• BMI (Body Mass Index) calculator\n"
        "• Interactive analytics dashboard with multiple chart types\n"
        "• Prediction history tracking and management\n"
        "• Data export functionality (CSV format)\n"
        "• Responsive web design for multiple devices\n\n"
        "The system is intended for educational and informational purposes and should not replace "
        "professional medical diagnosis."
    )
    scope.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()
    
    # ==================== CHAPTER 2: LITERATURE REVIEW ====================
    doc.add_heading("CHAPTER 2: LITERATURE REVIEW", level=1)
    
    doc.add_heading("2.1 Existing Systems", level=2)
    
    existing = doc.add_paragraph()
    existing.add_run(
        "Several healthcare prediction systems exist in the market, each with their own strengths "
        "and limitations:\n\n"
        "1. Hospital-based Diagnostic Systems\n"
        "   - Require expensive equipment and trained personnel\n"
        "   - Limited accessibility for general public\n"
        "   - High cost of operation\n\n"
        "2. Mobile Health Applications\n"
        "   - Often focus on single health metrics (e.g., BMI only)\n"
        "   - Limited prediction capabilities\n"
        "   - May lack scientific validation\n\n"
        "3. Online Health Calculators\n"
        "   - Basic risk calculators without machine learning\n"
        "   - No personalized recommendations\n"
        "   - Cannot track historical data\n\n"
        "4. Research-oriented ML Systems\n"
        "   - Complex interfaces not suitable for end users\n"
        "   - Require technical expertise to operate\n"
        "   - Limited to laboratory settings"
    )
    existing.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading("2.2 Proposed System", level=2)
    
    proposed = doc.add_paragraph()
    proposed.add_run(
        "The proposed Healthcare Data Analysis system addresses the limitations of existing systems "
        "by providing:\n\n"
        "• A user-friendly web interface accessible from any device\n"
        "• Machine learning-based predictions using validated datasets\n"
        "• Multi-disease assessment capability\n"
        "• Comprehensive visualization dashboards\n"
        "• Historical tracking of health assessments\n"
        "• Personalized health recommendations\n"
        "• Secure authentication and data protection\n"
        "• Export functionality for data portability\n\n"
        "The system uses Flask as the backend framework with MongoDB for data storage, ensuring "
        "scalability and flexibility. Machine learning models are trained on publicly available, "
        "validated medical datasets from UCI Machine Learning Repository."
    )
    proposed.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading("2.3 Advantages of Proposed System", level=2)
    
    adv = doc.add_paragraph()
    adv.add_run(
        "Key advantages of the proposed system:\n\n"
        "1. Accessibility\n"
        "   - Web-based platform accessible from anywhere\n"
        "   - No specialized hardware required\n"
        "   - Free to use for personal health monitoring\n\n"
        "2. Accuracy\n"
        "   - Uses ensemble machine learning algorithms\n"
        "   - Trained on validated medical datasets\n"
        "   - Cross-validated for reliable predictions\n\n"
        "3. Comprehensiveness\n"
        "   - Multiple disease predictions\n"
        "   - Holistic health score calculation\n"
        "   - Detailed analytics and visualizations\n\n"
        "4. Security\n"
        "   - CSRF protection on all forms\n"
        "   - Rate-limited login attempts\n"
        "   - Encrypted password storage\n"
        "   - Session timeout for inactive users\n\n"
        "5. User Experience\n"
        "   - Intuitive interface design\n"
        "   - Dark mode support\n"
        "   - Mobile-responsive layout"
    )
    adv.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()
    
    # ==================== CHAPTER 3: SYSTEM ANALYSIS ====================
    doc.add_heading("CHAPTER 3: SYSTEM ANALYSIS", level=1)
    
    doc.add_heading("3.1 Feasibility Study", level=2)
    
    feas = doc.add_paragraph()
    feas.add_run("Technical Feasibility\n").bold = True
    feas.add_run(
        "The project uses well-established technologies:\n"
        "• Python - Widely used for machine learning applications\n"
        "• Flask - Lightweight and flexible web framework\n"
        "• MongoDB - Scalable NoSQL database\n"
        "• scikit-learn - Industry-standard ML library\n\n"
    )
    feas.add_run("Economic Feasibility\n").bold = True
    feas.add_run(
        "All technologies used are open-source and free:\n"
        "• No licensing costs\n"
        "• Free tier available for MongoDB Atlas\n"
        "• Minimal hosting requirements\n\n"
    )
    feas.add_run("Operational Feasibility\n").bold = True
    feas.add_run(
        "The system is designed for ease of use:\n"
        "• Intuitive web interface\n"
        "• No training required for end users\n"
        "• Automated model predictions\n"
        "• Self-explanatory results with health tips"
    )
    feas.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading("3.2 System Requirements", level=2)
    
    req = doc.add_paragraph()
    req.add_run("Functional Requirements:\n\n")
    req.add_run(
        "1. User Management\n"
        "   • User registration with email verification\n"
        "   • Secure login with password encryption\n"
        "   • Password reset via OTP\n"
        "   • Profile management\n\n"
        "2. Health Assessment\n"
        "   • Health data input form\n"
        "   • Multi-disease prediction\n"
        "   • Health score calculation\n"
        "   • Risk level classification\n\n"
        "3. Data Visualization\n"
        "   • Dashboard with summary statistics\n"
        "   • Analytics charts and graphs\n"
        "   • Trend analysis over time\n\n"
        "4. Data Management\n"
        "   • Prediction history storage\n"
        "   • CSV export functionality\n"
        "   • Report generation\n\n"
        "Non-Functional Requirements:\n\n"
        "• Performance: Response time < 3 seconds\n"
        "• Security: Encrypted data transmission\n"
        "• Availability: 99% uptime\n"
        "• Scalability: Support for concurrent users"
    )
    req.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading("3.3 Software Requirements", level=2)
    
    sw_headers = ["Component", "Requirement"]
    sw_data = [
        ["Operating System", "Windows 10/11, Linux, macOS"],
        ["Programming Language", "Python 3.8+"],
        ["Web Framework", "Flask 3.0.0"],
        ["Database", "MongoDB 6.0+"],
        ["ML Library", "scikit-learn 1.3.2"],
        ["Browser", "Chrome, Firefox, Edge, Safari"],
        ["IDE", "VS Code, PyCharm (optional)"],
    ]
    create_table_with_header(doc, sw_headers, sw_data)
    
    doc.add_paragraph()
    doc.add_heading("3.4 Hardware Requirements", level=2)
    
    hw_headers = ["Component", "Minimum", "Recommended"]
    hw_data = [
        ["Processor", "Intel Core i3 / AMD Ryzen 3", "Intel Core i5 / AMD Ryzen 5"],
        ["RAM", "4 GB", "8 GB"],
        ["Storage", "500 MB free space", "1 GB free space"],
        ["Network", "Internet connection", "Broadband connection"],
    ]
    create_table_with_header(doc, hw_headers, hw_data)
    
    doc.add_page_break()
    
    # ==================== CHAPTER 4: SYSTEM DESIGN ====================
    doc.add_heading("CHAPTER 4: SYSTEM DESIGN", level=1)
    
    doc.add_heading("4.1 System Architecture", level=2)
    
    arch = doc.add_paragraph()
    arch.add_run(
        "The Healthcare Data Analysis system follows a three-tier architecture:\n\n"
        "1. Presentation Layer (Frontend)\n"
        "   • HTML5 templates with Jinja2 templating\n"
        "   • CSS3 with custom styling and dark mode\n"
        "   • JavaScript for interactive features\n"
        "   • Chart.js for data visualizations\n\n"
        "2. Application Layer (Backend)\n"
        "   • Flask web server\n"
        "   • Route handlers for HTTP requests\n"
        "   • Session management\n"
        "   • ML model inference engine\n"
        "   • Business logic processing\n\n"
        "3. Data Layer\n"
        "   • MongoDB Atlas (cloud database)\n"
        "   • User data storage\n"
        "   • Prediction history\n"
        "   • Session data\n\n"
        "[Figure 4.1: System Architecture Diagram - Insert diagram here]"
    )
    arch.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading("4.2 Data Flow Diagrams", level=2)
    
    dfd = doc.add_paragraph()
    dfd.add_run("Context Level DFD (Level 0)\n\n").bold = True
    dfd.add_run(
        "The context level DFD shows the system as a single process with external entities:\n\n"
        "External Entities:\n"
        "• User - Interacts with the system\n"
        "• Database - Stores and retrieves data\n"
        "• ML Models - Provides predictions\n\n"
        "Data Flows:\n"
        "• User inputs health data\n"
        "• System returns predictions and visualizations\n"
        "• Database stores/retrieves user and prediction data\n\n"
        "[Figure 4.2: Context Level DFD - Insert diagram here]\n\n"
    )
    dfd.add_run("Level 1 DFD\n\n").bold = True
    dfd.add_run(
        "The Level 1 DFD decomposes the system into major processes:\n\n"
        "1.0 User Authentication\n"
        "   • Register, Login, Logout, Password Reset\n\n"
        "2.0 Health Assessment\n"
        "   • Input validation, Feature extraction, Prediction\n\n"
        "3.0 Data Visualization\n"
        "   • Dashboard generation, Chart rendering\n\n"
        "4.0 History Management\n"
        "   • Store predictions, Retrieve history, Export data\n\n"
        "[Figure 4.3: Level 1 DFD - Insert diagram here]"
    )
    dfd.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading("4.3 ER Diagram", level=2)
    
    er = doc.add_paragraph()
    er.add_run(
        "The Entity-Relationship diagram shows the database structure:\n\n"
        "Entities:\n"
        "• Users - Stores user account information\n"
        "• Predictions - Stores health prediction records\n"
        "• PasswordResets - Stores OTP for password recovery\n"
        "• LoginAttempts - Tracks failed login attempts\n\n"
        "Relationships:\n"
        "• User HAS MANY Predictions (1:N)\n"
        "• User HAS ONE PasswordReset (1:1)\n"
        "• User HAS MANY LoginAttempts (1:N)\n\n"
        "[Figure 4.4: ER Diagram - Insert diagram here]"
    )
    er.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading("4.4 Database Design", level=2)
    
    db = doc.add_paragraph()
    db.add_run("Users Collection Schema\n\n").bold = True
    
    users_headers = ["Field", "Type", "Description"]
    users_data = [
        ["_id", "ObjectId", "Unique identifier (auto-generated)"],
        ["username", "String", "Unique username"],
        ["email", "String", "Unique email address"],
        ["password", "Binary", "Bcrypt hashed password"],
        ["created_at", "DateTime", "Account creation timestamp"],
        ["last_login", "DateTime", "Last login timestamp"],
        ["login_count", "Number", "Total successful logins"],
        ["is_active", "Boolean", "Account status"],
    ]
    create_table_with_header(doc, users_headers, users_data)
    
    doc.add_paragraph()
    
    pred_para = doc.add_paragraph()
    pred_para.add_run("Predictions Collection Schema\n\n").bold = True
    
    pred_headers = ["Field", "Type", "Description"]
    pred_data = [
        ["_id", "ObjectId", "Unique identifier"],
        ["user_id", "String", "Reference to user"],
        ["disease_type", "String", "'heart' or 'diabetes'"],
        ["input_data", "Object", "Health parameters"],
        ["result", "String", "'High Risk' or 'Low Risk'"],
        ["probability", "Number", "Risk percentage (0-100)"],
        ["health_score", "Number", "Overall health score (0-100)"],
        ["tips", "Array", "Personalized health tips"],
        ["timestamp", "DateTime", "Prediction timestamp"],
    ]
    create_table_with_header(doc, pred_headers, pred_data)
    
    doc.add_page_break()
    
    # ==================== CHAPTER 5: IMPLEMENTATION ====================
    doc.add_heading("CHAPTER 5: IMPLEMENTATION", level=1)
    
    doc.add_heading("5.1 Technology Stack", level=2)
    
    tech_headers = ["Layer", "Technology", "Version", "Purpose"]
    tech_data = [
        ["Backend", "Python", "3.8+", "Primary programming language"],
        ["Backend", "Flask", "3.0.0", "Web application framework"],
        ["Database", "MongoDB Atlas", "6.0+", "Cloud NoSQL database"],
        ["Database", "PyMongo", "4.6.1", "MongoDB Python driver"],
        ["ML", "scikit-learn", "1.3.2", "Machine learning library"],
        ["ML", "NumPy", "1.24.4", "Numerical computations"],
        ["ML", "Pandas", "2.0.3", "Data manipulation"],
        ["ML", "imbalanced-learn", "0.11.0", "SMOTE oversampling"],
        ["Frontend", "HTML5", "-", "Page structure"],
        ["Frontend", "CSS3", "-", "Styling and layout"],
        ["Frontend", "JavaScript", "ES6", "Interactive features"],
        ["Frontend", "Chart.js", "4.x", "Data visualization"],
        ["Security", "bcrypt", "4.1.2", "Password hashing"],
        ["Security", "certifi", "2023.11", "SSL certificates"],
    ]
    create_table_with_header(doc, tech_headers, tech_data)
    
    doc.add_paragraph()
    doc.add_heading("5.2 Machine Learning Models", level=2)
    
    ml = doc.add_paragraph()
    ml.add_run("Heart Disease Model\n\n").bold = True
    
    heart_headers = ["Property", "Value"]
    heart_data = [
        ["Algorithm", "Random Forest Classifier"],
        ["Data Source", "UCI Cleveland Heart Disease Database"],
        ["Original Samples", "297 real patient records"],
        ["Training Samples", "1500 (augmented with SMOTE)"],
        ["Accuracy", "~81% (5-fold cross-validation)"],
        ["AUC-ROC", "~99%"],
        ["Number of Trees", "100"],
        ["Max Depth", "10"],
    ]
    create_table_with_header(doc, heart_headers, heart_data)
    
    doc.add_paragraph()
    
    diab_para = doc.add_paragraph()
    diab_para.add_run("Diabetes Model\n\n").bold = True
    
    diab_headers = ["Property", "Value"]
    diab_data = [
        ["Algorithm", "Gradient Boosting Classifier"],
        ["Data Source", "Pima Indians Diabetes Database"],
        ["Original Samples", "768 real patient records"],
        ["Training Samples", "1500 (augmented with SMOTE)"],
        ["Accuracy", "~75% (5-fold cross-validation)"],
        ["AUC-ROC", "~94%"],
        ["Learning Rate", "0.1"],
        ["Number of Estimators", "100"],
    ]
    create_table_with_header(doc, diab_headers, diab_data)
    
    doc.add_paragraph()
    
    feat_para = doc.add_paragraph()
    feat_para.add_run("Input Features (16 features)\n\n").bold = True
    
    feat_headers = ["#", "Feature", "Description", "Range"]
    feat_data = [
        ["1", "age", "Patient age in years", "25-80"],
        ["2", "sex", "Gender (0=Female, 1=Male)", "0-1"],
        ["3", "trestbps", "Resting blood pressure (mm Hg)", "80-200"],
        ["4", "chol", "Serum cholesterol (mg/dl)", "100-400"],
        ["5", "fbs", "Fasting blood sugar > 120 mg/dl", "0-1"],
        ["6", "restecg", "Resting ECG results", "0-2"],
        ["7", "thalach", "Maximum heart rate achieved", "70-210"],
        ["8", "exang", "Exercise induced angina", "0-1"],
        ["9", "oldpeak", "ST depression induced by exercise", "0-6"],
        ["10", "slope", "Slope of peak exercise ST segment", "0-2"],
        ["11", "ca", "Number of major vessels colored", "0-3"],
        ["12", "thal", "Thalassemia", "1-3"],
        ["13", "smoking", "Patient smokes", "0-1"],
        ["14", "exercise", "Regular exercise", "0-1"],
        ["15", "alcohol", "Alcohol consumption", "0-1"],
        ["16", "bmi", "Body Mass Index", "15-50"],
    ]
    create_table_with_header(doc, feat_headers, feat_data)
    
    doc.add_page_break()
    
    doc.add_heading("5.3 Module Description", level=2)
    
    mod = doc.add_paragraph()
    mod.add_run("1. User Authentication Module\n\n").bold = True
    mod.add_run(
        "Handles all user-related operations:\n"
        "• Registration with email uniqueness validation\n"
        "• Login with rate limiting (5 attempts / 15 minutes)\n"
        "• Password reset via email OTP\n"
        "• Session management with 30-minute timeout\n"
        "• CSRF protection on all POST requests\n\n"
    )
    mod.add_run("2. Prediction Module\n\n").bold = True
    mod.add_run(
        "Core ML prediction functionality:\n"
        "• Input validation and preprocessing\n"
        "• Feature extraction and normalization\n"
        "• Model inference using joblib\n"
        "• Probability calculation\n"
        "• Health score computation\n"
        "• Personalized tips generation\n\n"
    )
    mod.add_run("3. Visualization Module\n\n").bold = True
    mod.add_run(
        "Data visualization components:\n"
        "• Dashboard summary cards\n"
        "• Risk distribution pie charts\n"
        "• Timeline charts for trends\n"
        "• BMI, blood pressure, sugar level charts\n"
        "• Health score progression charts\n\n"
    )
    mod.add_run("4. History Management Module\n\n").bold = True
    mod.add_run(
        "Prediction history operations:\n"
        "• Store predictions in MongoDB\n"
        "• Retrieve paginated history\n"
        "• Search and filter functionality\n"
        "• CSV export for all predictions\n"
        "• Individual report download\n\n"
    )
    mod.add_run("5. API Module\n\n").bold = True
    mod.add_run(
        "RESTful API endpoints:\n"
        "• /api/chart-data - Basic chart data\n"
        "• /api/analytics-data - Full analytics data\n"
        "• /api/health-score - Health score calculation\n"
    )
    mod.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading("5.4 Code Implementation", level=2)
    
    code = doc.add_paragraph()
    code.add_run("Key Code Snippets\n\n").bold = True
    code.add_run(
        "1. Flask Application Setup (app.py)\n\n"
        "app = Flask(__name__)\n"
        "app.secret_key = os.environ.get('SECRET_KEY', 'hc-analysis-prod-key-2026')\n"
        "app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(minutes=30)\n\n"
        "2. MongoDB Connection\n\n"
        "client = MongoClient(MONGO_URI, tls=True, tlsCAFile=certifi.where())\n"
        "db = client['healthcare_db']\n"
        "users_col = db['users']\n"
        "predictions_col = db['predictions']\n\n"
        "3. Prediction Route\n\n"
        "@app.route('/predict', methods=['POST'])\n"
        "@login_required\n"
        "def predict():\n"
        "    # Extract and validate input\n"
        "    # Prepare features for model\n"
        "    # Make prediction\n"
        "    # Calculate health score\n"
        "    # Store in database\n"
        "    # Return results\n\n"
        "4. ML Model Loading\n\n"
        "heart_model = joblib.load('model/heart_model.pkl')\n"
        "diabetes_model = joblib.load('model/diabetes_model.pkl')\n\n"
        "5. Password Hashing\n\n"
        "hashed = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt())\n"
        "# Verification\n"
        "bcrypt.checkpw(password.encode('utf-8'), stored_hash)\n"
    )
    
    doc.add_page_break()
    
    # ==================== CHAPTER 6: TESTING ====================
    doc.add_heading("CHAPTER 6: TESTING", level=1)
    
    doc.add_heading("6.1 Testing Methodology", level=2)
    
    test_method = doc.add_paragraph()
    test_method.add_run(
        "The application was tested using multiple testing approaches:\n\n"
        "1. Unit Testing\n"
        "   • Individual function testing\n"
        "   • Model prediction accuracy testing\n"
        "   • Input validation testing\n\n"
        "2. Integration Testing\n"
        "   • Database connectivity testing\n"
        "   • API endpoint testing\n"
        "   • Session management testing\n\n"
        "3. System Testing\n"
        "   • End-to-end workflow testing\n"
        "   • User interface testing\n"
        "   • Cross-browser compatibility testing\n\n"
        "4. Security Testing\n"
        "   • CSRF token validation\n"
        "   • Rate limiting verification\n"
        "   • Password encryption verification\n\n"
        "5. Performance Testing\n"
        "   • Response time measurement\n"
        "   • Concurrent user simulation\n"
        "   • Database query optimization"
    )
    test_method.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading("6.2 Test Cases", level=2)
    
    test_headers = ["ID", "Test Case", "Input", "Expected Output", "Status"]
    test_data = [
        ["TC01", "User Registration", "Valid user data", "Account created", "Pass"],
        ["TC02", "Duplicate Email", "Existing email", "Error message", "Pass"],
        ["TC03", "Weak Password", "Short password", "Validation error", "Pass"],
        ["TC04", "Valid Login", "Correct credentials", "Dashboard redirect", "Pass"],
        ["TC05", "Invalid Login", "Wrong password", "Error + attempts count", "Pass"],
        ["TC06", "Rate Limiting", "6 failed attempts", "Account locked", "Pass"],
        ["TC07", "Health Prediction", "Valid health data", "Risk prediction", "Pass"],
        ["TC08", "Missing Fields", "Incomplete form", "Validation error", "Pass"],
        ["TC09", "CSV Export", "Export request", "CSV file download", "Pass"],
        ["TC10", "Session Timeout", "30 min inactivity", "Auto logout", "Pass"],
        ["TC11", "CSRF Protection", "Invalid token", "Request rejected", "Pass"],
        ["TC12", "Password Reset", "Valid email + OTP", "Password changed", "Pass"],
    ]
    create_table_with_header(doc, test_headers, test_data)
    
    doc.add_heading("6.3 Results", level=2)
    
    results = doc.add_paragraph()
    results.add_run(
        "Testing Summary:\n\n"
        "• Total Test Cases: 12\n"
        "• Passed: 12\n"
        "• Failed: 0\n"
        "• Pass Rate: 100%\n\n"
        "All critical functionality has been tested and verified to work as expected. "
        "The system handles edge cases appropriately and provides meaningful error messages."
    )
    results.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()
    
    # ==================== CHAPTER 7: RESULTS AND DISCUSSION ====================
    doc.add_heading("CHAPTER 7: RESULTS AND DISCUSSION", level=1)
    
    doc.add_heading("7.1 Screenshots", level=2)
    
    screenshots = doc.add_paragraph()
    screenshots.add_run(
        "[Insert screenshots of the application here]\n\n"
        "Figure 7.1: Home Page\n"
        "- Modern landing page with gradient background\n"
        "- Navigation menu with Login/Register options\n"
        "- Security badges and feature highlights\n\n"
        "Figure 7.2: Registration Page\n"
        "- Clean form design with validation\n"
        "- Password strength indicator\n"
        "- CSRF token hidden field\n\n"
        "Figure 7.3: Login Page\n"
        "- Username and password fields\n"
        "- Rate limiting notice\n"
        "- Password reset link\n\n"
        "Figure 7.4: Dashboard\n"
        "- Summary statistics cards\n"
        "- Risk distribution chart\n"
        "- Recent predictions timeline\n"
        "- Quick action buttons\n\n"
        "Figure 7.5: Health Assessment Form\n"
        "- Disease type selector\n"
        "- Multiple input sections\n"
        "- Lifestyle factors checkboxes\n"
        "- Submit button with loading indicator\n\n"
        "Figure 7.6: Prediction Result Page\n"
        "- Risk level indicator (High/Low)\n"
        "- Probability percentage\n"
        "- Health score gauge\n"
        "- Personalized health tips\n\n"
        "Figure 7.7: Analytics Dashboard\n"
        "- Multiple chart types\n"
        "- Risk distribution analysis\n"
        "- Trend visualizations\n"
        "- Health metrics over time\n\n"
        "Figure 7.8: Prediction History\n"
        "- Tabular data with pagination\n"
        "- Search and filter options\n"
        "- Export buttons\n"
        "- Individual report download"
    )
    screenshots.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading("7.2 Model Performance", level=2)
    
    perf_headers = ["Metric", "Heart Disease Model", "Diabetes Model"]
    perf_data = [
        ["Accuracy", "81%", "75%"],
        ["Precision", "79%", "73%"],
        ["Recall", "83%", "77%"],
        ["F1-Score", "81%", "75%"],
        ["AUC-ROC", "99%", "94%"],
    ]
    create_table_with_header(doc, perf_headers, perf_data)
    
    doc.add_paragraph()
    
    perf_discussion = doc.add_paragraph()
    perf_discussion.add_run(
        "The machine learning models demonstrate good performance for health risk prediction:\n\n"
        "• Heart Disease Model achieves 81% accuracy using Random Forest, which is consistent "
        "with literature benchmarks for the UCI Cleveland dataset.\n\n"
        "• Diabetes Model achieves 75% accuracy using Gradient Boosting, which is reasonable "
        "given the inherent complexity of diabetes prediction.\n\n"
        "• High AUC-ROC scores indicate excellent discrimination ability between high-risk "
        "and low-risk cases.\n\n"
        "• SMOTE oversampling helped balance the datasets and improve model generalization."
    )
    perf_discussion.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()
    
    # ==================== CHAPTER 8: CONCLUSION ====================
    doc.add_heading("CHAPTER 8: CONCLUSION AND FUTURE SCOPE", level=1)
    
    doc.add_heading("8.1 Conclusion", level=2)
    
    conclusion = doc.add_paragraph()
    conclusion.add_run(
        "The Healthcare Data Analysis project successfully demonstrates the application of "
        "machine learning and data visualization techniques for health risk prediction. "
        "The system provides a comprehensive platform that:\n\n"
        "• Predicts heart disease and diabetes risk with reasonable accuracy\n"
        "• Provides interactive visualizations for health data analysis\n"
        "• Offers a secure, user-friendly web interface\n"
        "• Enables tracking of health assessments over time\n"
        "• Generates personalized health recommendations\n\n"
        "The project achieves all its stated objectives and provides a solid foundation "
        "for healthcare data analysis applications. The use of ensemble machine learning "
        "algorithms (Random Forest and Gradient Boosting) ensures reliable predictions, "
        "while the Flask-MongoDB architecture provides scalability and flexibility.\n\n"
        "The security features including CSRF protection, rate limiting, and encrypted "
        "password storage ensure user data protection. The responsive design makes the "
        "application accessible across different devices."
    )
    conclusion.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_heading("8.2 Future Enhancements", level=2)
    
    future = doc.add_paragraph()
    future.add_run(
        "The following enhancements can be implemented in future versions:\n\n"
        "1. Additional Disease Predictions\n"
        "   • Hypertension risk assessment\n"
        "   • Kidney disease prediction\n"
        "   • Cancer risk screening\n\n"
        "2. Advanced Features\n"
        "   • Integration with wearable devices\n"
        "   • Real-time health monitoring\n"
        "   • Telemedicine consultation booking\n"
        "   • Medical report upload and analysis\n\n"
        "3. AI Improvements\n"
        "   • Deep learning models for improved accuracy\n"
        "   • Natural language processing for symptom analysis\n"
        "   • Explainable AI for better interpretability\n\n"
        "4. User Experience\n"
        "   • Mobile application development\n"
        "   • Multi-language support\n"
        "   • Voice-based interaction\n"
        "   • Personalized health plans\n\n"
        "5. Healthcare Integration\n"
        "   • Electronic Health Records (EHR) integration\n"
        "   • Doctor referral system\n"
        "   • Insurance claim assistance\n"
        "   • Medication reminders"
    )
    future.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()
    
    # ==================== REFERENCES ====================
    ref_title = doc.add_paragraph()
    ref_title.add_run("REFERENCES").bold = True
    ref_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref_title.runs[0].font.size = Pt(16)
    
    doc.add_paragraph()
    
    refs = doc.add_paragraph()
    refs.add_run(
        "[1] Dua, D. and Graff, C. (2019). UCI Machine Learning Repository. "
        "Irvine, CA: University of California, School of Information and Computer Science. "
        "http://archive.ics.uci.edu/ml\n\n"
        "[2] Smith, J.W., Everhart, J.E., Dickson, W.C., Knowler, W.C., & Johannes, R.S. (1988). "
        "Using the ADAP learning algorithm to forecast the onset of diabetes mellitus. "
        "Proceedings of the Symposium on Computer Applications and Medical Care, 261-265.\n\n"
        "[3] Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. "
        "Journal of Machine Learning Research, 12, 2825-2830.\n\n"
        "[4] Chawla, N.V., Bowyer, K.W., Hall, L.O., & Kegelmeyer, W.P. (2002). "
        "SMOTE: Synthetic Minority Over-sampling Technique. "
        "Journal of Artificial Intelligence Research, 16, 321-357.\n\n"
        "[5] Flask Documentation. (2024). Flask: A Python Microframework. "
        "https://flask.palletsprojects.com/\n\n"
        "[6] MongoDB Documentation. (2024). MongoDB Atlas Database. "
        "https://www.mongodb.com/docs/\n\n"
        "[7] Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5-32.\n\n"
        "[8] Friedman, J.H. (2001). Greedy Function Approximation: A Gradient Boosting Machine. "
        "Annals of Statistics, 29(5), 1189-1232.\n\n"
        "[9] Chart.js Documentation. (2024). Simple yet flexible JavaScript charting. "
        "https://www.chartjs.org/\n\n"
        "[10] OWASP. (2024). Cross-Site Request Forgery Prevention Cheat Sheet. "
        "https://cheatsheetseries.owasp.org/cheatsheets/Cross-Site_Request_Forgery_Prevention_Cheat_Sheet.html"
    )
    refs.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()
    
    # ==================== APPENDIX ====================
    app_title = doc.add_paragraph()
    app_title.add_run("APPENDIX").bold = True
    app_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    app_title.runs[0].font.size = Pt(16)
    
    doc.add_paragraph()
    
    doc.add_heading("A. Project File Structure", level=2)
    
    structure = doc.add_paragraph()
    structure.add_run(
        "Healthcare Data Analysis/\n"
        "├── app.py                    # Main Flask application (1000+ lines)\n"
        "├── train_model.py            # ML model training script\n"
        "├── requirements.txt          # Python dependencies\n"
        "├── README.md                 # Basic readme\n"
        "├── DOCUMENTATION.md          # Detailed documentation\n"
        "│\n"
        "├── model/                    # Trained ML models\n"
        "│   ├── heart_model.pkl       # Heart disease classifier\n"
        "│   ├── diabetes_model.pkl    # Diabetes risk classifier\n"
        "│   └── model_meta.pkl        # Model metadata\n"
        "│\n"
        "├── static/\n"
        "│   ├── css/\n"
        "│   │   └── style.css         # Application styles (1000+ lines)\n"
        "│   └── js/\n"
        "│       └── main.js           # Frontend JavaScript (500+ lines)\n"
        "│\n"
        "└── templates/                # HTML templates (12 files)\n"
        "    ├── home.html\n"
        "    ├── login.html\n"
        "    ├── register.html\n"
        "    ├── dashboard.html\n"
        "    ├── health_form.html\n"
        "    ├── result.html\n"
        "    ├── history.html\n"
        "    ├── analytics.html\n"
        "    ├── bmi_calculator.html\n"
        "    ├── profile.html\n"
        "    ├── forgot_password.html\n"
        "    ├── verify_otp.html\n"
        "    └── reset_password.html\n"
    )
    
    doc.add_heading("B. Requirements.txt", level=2)
    
    requirements = doc.add_paragraph()
    requirements.add_run(
        "Flask==3.0.0\n"
        "pymongo==4.6.1\n"
        "bcrypt==4.1.2\n"
        "scikit-learn==1.3.2\n"
        "pandas==2.0.3\n"
        "numpy==1.24.4\n"
        "joblib==1.3.2\n"
        "certifi==2023.11.17\n"
        "imbalanced-learn==0.11.0\n"
    )
    
    doc.add_heading("C. Running the Application", level=2)
    
    running = doc.add_paragraph()
    running.add_run(
        "Step 1: Install Python 3.8 or higher\n\n"
        "Step 2: Create virtual environment (optional)\n"
        "        python -m venv venv\n"
        "        venv\\Scripts\\activate  (Windows)\n\n"
        "Step 3: Install dependencies\n"
        "        pip install -r requirements.txt\n\n"
        "Step 4: Train models (if not present)\n"
        "        python train_model.py\n\n"
        "Step 5: Run the application\n"
        "        python app.py\n\n"
        "Step 6: Open browser\n"
        "        http://localhost:5000\n"
    )
    
    # Save document
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "Healthcare_Data_Analysis_Project_Documentation.docx"
    )
    doc.save(output_path)
    print(f"✅ Documentation saved to: {output_path}")
    return output_path

if __name__ == "__main__":
    create_documentation()
